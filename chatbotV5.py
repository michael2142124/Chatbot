import streamlit as st
import PyPDF2
import docx
import os
import pandas as pd
import requests
from bs4 import BeautifulSoup
from readability import Document
from dotenv import load_dotenv
from openai import OpenAI

# Load API key
load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

st.set_page_config(page_title="🧾 Document Chatbot", layout="wide")

# --- Styling ---
st.markdown("""
    <style>
        .message {
            padding: 1em;
            border-radius: 0.5em;
            margin-bottom: 0.8em;
            color: black;
        }
        .user {
            background-color: #f0f0f0;
        }
        .assistant {
            background-color: #e8f5e9;
        }
    </style>
""", unsafe_allow_html=True)

st.title("Risk Drive Chatbot")

# --- Load Documents from /docs ---
DOC_FOLDER = "docs"

def load_documents(folder_path):
    text = ""
    for root, _, files in os.walk(folder_path):
        for filename in files:
            file_path = os.path.join(root, filename)
            try:
                if filename.endswith(".pdf"):
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                elif filename.endswith(".docx"):
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                elif filename.endswith(".xlsx"):
                    xls = pd.ExcelFile(file_path)
                    for sheet in xls.sheet_names:
                        df = pd.read_excel(xls, sheet_name=sheet)
                        text += f"\n[Sheet: {sheet}]\n"
                        text += df.astype(str).to_string(index=False) + "\n"
            except Exception as e:
                print(f"❌ Skipped {filename}: {e}")
    return text

# --- Scrape Full Legal Text from Justice Canada ---
def fetch_justice_law_content(base_url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        base_resp = requests.get(base_url, headers=headers, timeout=10)
        base_resp.raise_for_status()

        base_soup = BeautifulSoup(base_resp.text, "html.parser")
        base_domain = "https://laws-lois.justice.gc.ca"

        section_links = set()
        for a in base_soup.find_all("a", href=True):
            href = a["href"]
            if href.startswith("/eng/acts/") or href.startswith("/eng/regulations/"):
                if "page-" in href or "FullText.html" in href:
                    full_url = base_domain + href
                    section_links.add(full_url)

        # If no sections found, fallback to base page
        if not section_links:
            section_links.add(base_url)

        full_text = ""
        for link in sorted(section_links):
            try:
                resp = requests.get(link, headers=headers, timeout=10)
                resp.raise_for_status()
                doc = Document(resp.text)
                title = doc.short_title()
                summary_html = doc.summary()
                soup = BeautifulSoup(summary_html, "html.parser")
                text = soup.get_text(separator="\n")
                cleaned = "\n".join([line.strip() for line in text.splitlines() if line.strip()])
                full_text += f"\n\n--- {title} ---\n{cleaned}"
            except Exception as e:
                print(f"❌ Failed to fetch section: {link}: {e}")
        return full_text[:20000]  # Limit to avoid token overflow
    except Exception as e:
        print(f"❌ Error fetching base page: {e}")
        return None

# --- Load All Context at Startup ---
@st.cache_data(show_spinner=False)
def load_all_context():
    doc_text = load_documents(DOC_FOLDER)

    LAW_URLS = [
        # TDGR - Transportation of Dangerous Goods Regulations
        "https://laws-lois.justice.gc.ca/eng/regulations/SOR-2002-184/index.html",

        # PIPEDA - Personal Information Protection and Electronic Documents Act
        "https://laws-lois.justice.gc.ca/eng/acts/P-8.6/",

        # PCMLTFA - Proceeds of Crime (Money Laundering) and Terrorist Financing Act
        "https://laws-lois.justice.gc.ca/eng/acts/P-24.501/index.html",

        # PCMLTFA Regulations (Reporting Obligations)
        "https://laws-lois.justice.gc.ca/eng/regulations/SOR-2001-317/index.html"
    ]

    law_text = ""
    for url in LAW_URLS:
        law_text += fetch_justice_law_content(url) or ""

    combined = doc_text + "\n" + law_text
    return combined[:16000]  # Token-safe limit

# Initialize content
if "document_text" not in st.session_state:
    with st.spinner("🔄 Loading documents and legal websites..."):
        st.session_state.document_text = load_all_context()

# --- Set up chat history ---
if "history" not in st.session_state:
    st.session_state.history = [
        {
            "role": "system",
            "content": "You are a helpful assistant. Only answer based on the documents and legal text provided."
        },
        {
            "role": "user",
            "content": f"Use the following context to answer questions:\n\n{st.session_state.document_text}"
        }
    ]

# --- Chat UI ---
st.markdown("### 💬 Ask a question about our documents")
chat_placeholder = st.container()

with chat_placeholder:
    for msg in st.session_state.history[2:]:
        if msg["role"] == "user":
            st.markdown(f"<div class='message user'>💬 <strong>You:</strong> {msg['content']}</div>", unsafe_allow_html=True)
        elif msg["role"] == "assistant":
            st.markdown(f"<div class='message assistant'>🤖 <strong>Chatbot:</strong> {msg['content']}</div>", unsafe_allow_html=True)

# --- Input Box ---
with st.form("chat_form", clear_on_submit=True):
    user_input = st.text_input("Your question:", placeholder="How can I help you?...", label_visibility="collapsed")
    submitted = st.form_submit_button("Send", use_container_width=True)

if submitted and user_input:
    st.session_state.history.append({"role": "user", "content": user_input})
    with st.spinner("🤔 Thinking..."):
        reply = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=st.session_state.history,
            temperature=0.2,
            max_tokens=500
        )
        response_text = reply.choices[0].message.content
        st.session_state.history.append({"role": "assistant", "content": response_text})
    st.rerun()
